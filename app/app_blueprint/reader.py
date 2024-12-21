from flask import Blueprint, render_template, request, jsonify, flash, redirect, url_for, abort
import sys
import os
import requests
import json
from flask import current_app
from openai import OpenAI

# 添加项目根目录到Python路径
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

from Database.model import (
    Documents, Users, Comments, Highlights,
    Notes, Evernote, AuditLog, Folders,
    DocKeywords, Contractors, Relations, Participants, People,
    FolderContents, Corrections
)
from Database.config import db
from flask_login import login_required, current_user
from datetime import datetime
from sqlalchemy import text
import logging
import traceback
from io import BytesIO
from docx import Document
from flask import send_file

# 配置日志
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

reader_bp = Blueprint('reader', __name__, url_prefix='/reader')

def validate_document_access(doc_id):
    """验证文档访问权限"""
    doc = Documents.query.get_or_404(doc_id)
    if not doc:
        abort(404, description="Document not found")
    return doc

def create_audit_log(user_id, action, description, table_name, record_id=None):
    """创建审计日志"""
    try:
        # 确保所有字段都有值
        if not all([user_id, action, description, table_name]):
            logger.error("审计日志缺少必要字段")
            return False

        # 创建审计日志记录
        audit_log = AuditLog(
            User_id=user_id,
            Audit_actionType=action,
            Audit_actionDescription=description,
            Audit_targetTable=table_name,
            Audit_timestamp=datetime.now()
        )
        
        db.session.add(audit_log)
        db.session.commit()
        
        return True
    except Exception as e:
        logger.error(f"创建审计日志失败: {str(e)}")
        db.session.rollback()
        return False

@reader_bp.route('/document/<doc_id>')
def document(doc_id):
    try:
        doc = validate_document_access(doc_id)
        
        # 获取关键词
        keywords = DocKeywords.query.filter_by(Doc_id=doc_id).all()
        
        # 获取契约人信息
        contractors_info = Contractors.query.filter_by(Doc_id=doc_id).first()
        contractors = []
        relation = None
        if contractors_info:
            alice = People.query.get(contractors_info.Alice_id)
            bob = People.query.get(contractors_info.Bob_id)
            if alice and bob:
                contractors = [
                    {"name": alice.Person_name},
                    {"name": bob.Person_name}
                ]
                # 获取关系
                relation = Relations.query.filter_by(
                    Alice_id=contractors_info.Alice_id,
                    Bob_id=contractors_info.Bob_id
                ).first()
                if relation:
                    relation = relation.Relation_type
        
        # 获取参人信息
        participants_query = db.session.query(
            Participants, People
        ).join(
            People, Participants.Person_id == People.Person_id
        ).filter(
            Participants.Doc_id == doc_id
        ).all()
        
        participants = [
            {"name": person.Person_name, "role": participant.Part_role}
            for participant, person in participants_query
        ]
        
        data = {
            'document': doc,
            'content': doc.Doc_originalText,
            'keywords': [{"KeyWord": kw.KeyWord} for kw in keywords],
            'contractors': contractors,
            'relation': relation,
            'participants': participants,
            'highlights': [],
            'notes': [],
            'comments': [],
            'evernotes': [],
            'corrections': []  # 添加纠错记录列表
        }
        
        if current_user.is_authenticated:
            # 获取用户相关数据
            data['highlights'] = Highlights.query.filter_by(
                Doc_id=doc_id,
                User_id=current_user.User_id
            ).all()
            
            # 获取批注
            data['notes'] = [{
                'id': note.Note_id,
                'content': note.Note_annotationText,
                'created_at': note.Note_createdAt.strftime('%Y-%m-%d %H:%M:%S')
            } for note in Notes.query.filter_by(
                Doc_id=doc_id,
                User_id=current_user.User_id
            ).all()]
            
            # 获取笔记
            evernotes = Evernote.query.filter_by(
                Doc_id=doc_id,
                User_id=current_user.User_id
            ).order_by(Evernote.Evernote_viewedAt.desc()).all()
            
            data['evernotes'] = [{
                'Evernote_id': note.Evernote_id,
                'Evernote_text': note.Evernote_text,
                'created_at': note.Evernote_viewedAt.strftime('%Y-%m-%d %H:%M:%S')
            } for note in evernotes]
            
            # 获取评论
            comments = Comments.query.filter_by(
                Doc_id=doc_id
            ).order_by(Comments.Comment_createdAt.desc()).all()
            
            data['comments'] = [{
                'id': comment.Comment_id,
                'content': comment.Comment_text,
                'user_name': Users.query.get(comment.User_id).User_name if comment.User_id else '匿名用户',
                'created_at': comment.Comment_createdAt.strftime('%Y-%m-%d %H:%M:%S'),
                'is_mine': comment.User_id == current_user.User_id
            } for comment in comments]
            
            # 获取错记录
            corrections = Corrections.query.filter_by(Doc_id=doc_id).all()
            data['corrections'] = [{
                'correction_id': c.Correction_id,
                'user_id': str(c.User_id),
                'user_name': Users.query.get(c.User_id).User_name if c.User_id else '系统',
                'text': c.Correction_text,
                'created_at': c.Correction_createdAt.strftime('%Y-%m-%d %H:%M:%S')
            } for c in corrections]
        
        return render_template('reader/document.html', **data)
        
    except Exception as e:
        logger.error(f"Error in document route: {str(e)}")
        flash('获取文档失败，请稍后再试', 'error')
        return redirect(url_for('searcher.index'))

@reader_bp.route('/add_highlight', methods=['POST'])
@login_required
def add_highlight():
    try:
        data = request.get_json()
        
        # 验证必要字段
        required_fields = ['doc_id', 'start_offset', 'end_offset', 'color']
        if not all(field in data for field in required_fields):
            return jsonify({
                'success': False,
                'error': 'Missing required fields'
            }), 400
            
        # 验证文档存在
        doc = validate_document_access(data['doc_id'])
        
        # 创建高亮
        new_highlight = Highlights(
            Doc_id=data['doc_id'],
            User_id=current_user.User_id,
            Highlight_startPosition=data['start_offset'],
            Highlight_endPosition=data['end_offset'],
            Highlight_color=data.get('color', '#ffeb3b'),  # 默认黄色
            Highlight_createdAt=datetime.now()
        )
        
        db.session.add(new_highlight)
        db.session.commit()
        
        # 只创建一条用户操作的审计日志
        create_audit_log(
            current_user.User_id,
            'INSERT',
            f'用户 {current_user.User_name} 在文档 {data["doc_id"]} 中添加了高亮',
            'Highlights',
            new_highlight.Highlight_id
        )
        
        return jsonify({
            'success': True,
            'data': {
                'id': new_highlight.Highlight_id,
                'start_offset': new_highlight.Highlight_startPosition,
                'end_offset': new_highlight.Highlight_endPosition,
                'color': new_highlight.Highlight_color
            }
        })
        
    except Exception as e:
        logger.error(f"Error adding highlight: {str(e)}")
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@reader_bp.route('/add_note', methods=['POST'])
@login_required
def add_note():
    try:
        data = request.get_json()
        
        # 验证必要字段
        required_fields = ['doc_id', 'content', 'start_offset', 'end_offset']
        if not all(field in data for field in required_fields):
            return jsonify({
                'success': False,
                'error': 'Missing required fields'
            }), 400
            
        # 验证文档存在
        doc = validate_document_access(data['doc_id'])
        
        # 创建批注
        new_note = Notes(
            Doc_id=data['doc_id'],
            User_id=current_user.User_id,
            Note_annotationText=data['content'],
            Note_startPosition=data['start_offset'],
            Note_endPosition=data['end_offset'],
            Note_createdAt=datetime.now()
        )
        
        db.session.add(new_note)
        db.session.commit()
        
        # 只创建一条用户操作的审计日志
        create_audit_log(
            current_user.User_id,
            'INSERT',
            f'用户 {current_user.User_name} 在文档 {data["doc_id"]} 中添加了批注',
            'Notes',
            new_note.Note_id
        )
        
        return jsonify({
            'success': True,
            'data': {
                'id': new_note.Note_id,
                'content': new_note.Note_annotationText,
                'created_at': new_note.Note_createdAt.strftime('%Y-%m-%d %H:%M:%S')
            }
        })
        
    except Exception as e:
        logger.error(f"Error adding note: {str(e)}")
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@reader_bp.route('/add_comment', methods=['POST'])
@login_required
def add_comment():
    try:
        data = request.get_json()
        
        # 验证必要字段
        if not data.get('Doc_id') or not data.get('Comment_text'):
            return jsonify({
                'success': False,
                'error': '缺少必要字段'
            }), 400
            
        # 创建评论
        comment = Comments(
            Doc_id=data['Doc_id'],
            User_id=current_user.User_id,
            Comment_text=data['Comment_text'],
            Comment_createdAt=datetime.now()
        )
        
        db.session.add(comment)
        db.session.commit()
        
        # 只创建一条用户操作的审计日志
        create_audit_log(
            current_user.User_id,
            'INSERT',
            f'用户 {current_user.User_name} 在文档 {data["Doc_id"]} 中添加了评论',
            'Comments',
            comment.Comment_id
        )
        
        # 返回新评论的完整信息
        return jsonify({
            'success': True,
            'data': {
                'id': comment.Comment_id,
                'content': comment.Comment_text,
                'user_name': current_user.User_name,
                'created_at': comment.Comment_createdAt.strftime('%Y-%m-%d %H:%M:%S'),
                'is_mine': True
            }
        })
        
    except Exception as e:
        logger.error(f"添加评论失败: {str(e)}")
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@reader_bp.route('/add_to_folder', methods=['POST'])
@login_required
def add_to_folder():
    try:
        data = request.get_json()
        
        # 验证必要字段
        if not data.get('doc_id') or not data.get('folder_name'):
            return jsonify({
                'success': False,
                'error': '缺少必要字段'
            }), 400
            
        # 找或创建收藏夹
        folder = Folders.query.filter_by(
            User_id=current_user.User_id,
            Folder_name=data['folder_name']
        ).first()
        
        if not folder:
            # 创建新收藏夹
            folder = Folders(
                User_id=current_user.User_id,
                Folder_name=data['folder_name'],
                Remarks=data.get('remarks'),  # 添加备注字段
                Folder_createdAt=datetime.now()
            )
            db.session.add(folder)
            db.session.commit()
        elif data.get('remarks'):  # 如果收藏夹已存在且提供了新的备注
            folder.Remarks = data['remarks']  # 更新备注
            db.session.commit()
        
        # 检查文档是否已在收藏夹中
        existing_content = FolderContents.query.filter_by(
            Folder_id=folder.Folder_id,
            Doc_id=data['doc_id']
        ).first()
        
        if existing_content:
            return jsonify({
                'success': False,
                'error': '文档已在该收藏夹中'
            }), 400
            
        # 创建新的收藏内容记录
        folder_content = FolderContents(
            Folder_id=folder.Folder_id,
            Doc_id=data['doc_id'],
            CreatedAt=datetime.now()
        )
        
        db.session.add(folder_content)
        db.session.commit()
        
        # 只创建一条用户操作的审计日志
        create_audit_log(
            current_user.User_id,
            'INSERT',
            f'用户 {current_user.User_name} 收藏了文档 {data["doc_id"]}',
            'FolderContents',
            folder_content.Content_id
        )
        
        return jsonify({
            'success': True,
            'data': {
                'id': folder.Folder_id,
                'folder_name': folder.Folder_name,
                'remarks': folder.Remarks
            }
        })
        
    except Exception as e:
        logger.error(f"添加到收藏夹失败: {str(e)}")
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@reader_bp.route('/get_highlights/<doc_id>')
@login_required
def get_highlights(doc_id):
    try:
        highlights = Highlights.query.filter_by(
            Doc_id=doc_id,
            User_id=current_user.User_id
        ).all()
        
        return jsonify({
            'success': True,
            'highlights': [{
                'Highlight_id': h.Highlight_id,
                'Highlight_startPosition': h.Highlight_startPosition,
                'Highlight_endPosition': h.Highlight_endPosition,
                'Highlight_color': h.Highlight_color
            } for h in highlights]
        })
        
    except Exception as e:
        logger.error(f"获取高亮列表失败: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@reader_bp.route('/get_notes/<doc_id>')
@login_required
def get_notes(doc_id):
    try:
        notes = Notes.query.filter_by(
            Doc_id=doc_id,
            User_id=current_user.User_id
        ).all()
        
        return jsonify({
            'success': True,
            'notes': [{
                'id': note.Note_id,
                'content': note.Note_annotationText,
                'start_offset': note.Note_startPosition,
                'end_offset': note.Note_endPosition,
                'created_at': note.Note_createdAt.strftime('%Y-%m-%d %H:%M:%S')
            } for note in notes]
        })
        
    except Exception as e:
        logger.error(f"获取批注失败: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@reader_bp.route('/get_comments/<doc_id>')
@login_required
def get_comments(doc_id):
    try:
        comments = Comments.query.filter_by(Doc_id=doc_id).order_by(Comments.Comment_createdAt.desc()).all()
        
        return jsonify({
            'success': True,
            'comments': [{
                'id': comment.Comment_id,
                'content': comment.Comment_text,
                'user_name': Users.query.get(comment.User_id).User_name if comment.User_id else '匿名用户',
                'created_at': comment.Comment_createdAt.strftime('%Y-%m-%d %H:%M:%S'),
                'is_mine': comment.User_id == current_user.User_id
            } for comment in comments]
        })
        
    except Exception as e:
        logger.error(f"获取评论失败: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@reader_bp.route('/export/<doc_id>/<text_type>')
@login_required
def export_document(doc_id, text_type):
    try:
        # 验证文档存在
        doc = validate_document_access(doc_id)
        
        # 创建Word文档
        word_doc = Document()
        
        # 添加标题
        word_doc.add_heading(doc.Doc_title, 0)
        
        # 添加基本信息
        word_doc.add_paragraph(f'文书类型: {doc.Doc_type}')
        if doc.Doc_summary:
            word_doc.add_paragraph(f'文书大意: {doc.Doc_summary}')
        
        # 添加正文
        word_doc.add_heading('正文', level=1)
        if text_type == 'simplified' and doc.Doc_simplifiedText:
            content = doc.Doc_simplifiedText
        else:
            content = doc.Doc_originalText
            
        word_doc.add_paragraph(content)
        
        # 添加用户的标注内容
        if current_user.is_authenticated:
            highlights = Highlights.query.filter_by(
                Doc_id=doc_id,
                User_id=current_user.User_id
            ).all()
            
            notes = Notes.query.filter_by(
                Doc_id=doc_id,
                User_id=current_user.User_id
            ).all()
            
            if highlights or notes:
                word_doc.add_heading('我的标注', level=1)
                
                if highlights:
                    word_doc.add_heading('高亮部分', level=2)
                    for h in highlights:
                        try:
                            highlighted_text = content[h.Highlight_startPosition:h.Highlight_endPosition]
                            p = word_doc.add_paragraph()
                            p.add_run(f'• {highlighted_text}')
                            p.add_run(f' (标注时间: {h.Highlight_createdAt.strftime("%Y-%m-%d %H:%M:%S")})')
                        except Exception as e:
                            logger.error(f"处理高��文本时出错: {str(e)}")
                
                if notes:
                    word_doc.add_heading('批注内容', level=2)
                    for n in notes:
                        try:
                            noted_text = content[n.Note_startPosition:n.Note_endPosition]
                            p = word_doc.add_paragraph()
                            p.add_run(f'原文: {noted_text}\n').bold = True
                            p.add_run(f'批注: {n.Note_annotationText}')
                            p.add_run(f'\n(批注时间: {n.Note_createdAt.strftime("%Y-%m-%d %H:%M:%S")})')
                        except Exception as e:
                            logger.error(f"处理批注文本时出错: {str(e)}")
        
        # 添加评论区
        comments = Comments.query.filter_by(Doc_id=doc_id).all()
        if comments:
            word_doc.add_heading('评论区', level=1)
            for comment in comments:
                p = word_doc.add_paragraph()
                p.add_run(f'• {comment.Comment_text}')
                p.add_run(f'\n  by {comment.user.User_name} ({comment.Comment_createdAt.strftime("%Y-%m-%d %H:%M:%S")})')
        
        # 保存到内存中
        file_stream = BytesIO()
        word_doc.save(file_stream)
        file_stream.seek(0)
        
        # 创建审��日志
        create_audit_log(
            current_user.User_id,
            'EXPORT',
            f'Exported document {doc_id} as {text_type}',
            'Documents',
            doc_id
        )
        
        # 返回文件
        filename = f"{doc.Doc_title}_{text_type}.docx"
        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        logger.error(f"导出文档时出错: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@reader_bp.route('/update_highlight', methods=['POST'])
@login_required
def update_highlight():
    try:
        data = request.get_json()
        
        # 验证必要字段
        required_fields = ['highlight_id', 'color']
        if not all(field in data for field in required_fields):
            return jsonify({
                'success': False,
                'error': 'Missing required fields'
            }), 400
        
        # 查找高亮记录
        highlight = Highlights.query.filter_by(
            Highlight_id=data['highlight_id'],
            User_id=current_user.User_id
        ).first()
        
        if not highlight:
            return jsonify({
                'success': False,
                'error': 'Highlight not found'
            }), 404
            
        # 更新颜色
        highlight.Highlight_color = data['color']
        db.session.commit()
        
        # 创建审计日志
        create_audit_log(
            current_user.User_id,
            'UPDATE',
            f'Updated highlight {data["highlight_id"]} color',
            'Highlights',
            highlight.Highlight_id
        )
        
        return jsonify({
            'success': True,
            'data': {
                'id': highlight.Highlight_id,
                'color': highlight.Highlight_color
            }
        })
        
    except Exception as e:
        logger.error(f"更新高亮颜色时出错: {str(e)}")
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@reader_bp.route('/delete_highlight/<int:highlight_id>', methods=['DELETE'])
@login_required
def delete_highlight(highlight_id):
    try:
        highlight = Highlights.query.filter_by(
            Highlight_id=highlight_id,
            User_id=current_user.User_id
        ).first()
        
        if not highlight:
            return jsonify({
                'success': False,
                'error': '高亮不存在或无权限删除'
            }), 404
            
        doc_id = highlight.Doc_id  # 保存文档ID用于审计日志
        db.session.delete(highlight)
        db.session.commit()
        
        # 修改审计日志格式
        create_audit_log(
            current_user.User_id,
            'DELETE',
            f'用户 {current_user.User_name} 删除了文档 {doc_id} 中的高亮',
            'Highlights',
            highlight_id
        )
        
        return jsonify({'success': True})
        
    except Exception as e:
        logger.error(f"删除高亮时出错: {str(e)}")
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@reader_bp.route('/delete_note/<int:note_id>', methods=['DELETE'])
@login_required
def delete_note(note_id):
    try:
        # 查找批注记录
        note = Notes.query.filter_by(
            Note_id=note_id,
            User_id=current_user.User_id
        ).first()
        
        if not note:
            return jsonify({
                'success': False,
                'error': '批注不存在或无权限删除'
            }), 404
            
        # 删除批注
        db.session.delete(note)
        db.session.commit()
        
        # 创建审计日志
        create_audit_log(
            user_id=current_user.User_id,
            action='DELETE',
            description=f'用户 {current_user.User_name} 删除了批注 {note_id}',
            table_name='Notes',
            record_id=note_id
        )
        
        return jsonify({
            'success': True,
            'message': '批注删除成功'
        })
        
    except Exception as e:
        logger.error(f"删除批注时出错: {str(e)}")
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@reader_bp.route('/remove_from_folder', methods=['POST'])
@login_required
def remove_from_folder():
    try:
        data = request.get_json()
        
        # 验证必要字段
        if not data.get('doc_id') or not data.get('folder_name'):
            return jsonify({
                'success': False,
                'error': 'Missing required fields'
            }), 400
            
        # 查找收藏夹
        folder = Folders.query.filter_by(
            User_id=current_user.User_id,
            Folder_name=data['folder_name']
        ).first()
        
        if not folder:
            return jsonify({
                'success': False,
                'error': '收藏夹不存在'
            }), 404
            
        # 查找收藏内容记录
        folder_content = FolderContents.query.filter_by(
            Folder_id=folder.Folder_id,
            Doc_id=data['doc_id']
        ).first()
        
        if not folder_content:
            return jsonify({
                'success': False,
                'error': '文档不在该收藏夹中'
            }), 404
            
        # 删除收藏内容记录
        db.session.delete(folder_content)
        db.session.commit()
        
        # 创建审计日志
        create_audit_log(
            current_user.User_id,
            'DELETE',
            f'从收藏夹 {data["folder_name"]} 中移除文档 {data["doc_id"]}',
            'FolderContents',
            folder_content.Content_id
        )
        
        return jsonify({
            'success': True,
            'message': '文档已从收藏夹中移除'
        })
        
    except Exception as e:
        logger.error(f"从收藏夹移除出错: {str(e)}")
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@reader_bp.route('/check_folder/<doc_id>/<folder_name>')
@login_required
def check_folder(doc_id, folder_name):
    try:
        # 查找收藏夹
        folder = Folders.query.filter_by(
            User_id=current_user.User_id,
            Folder_name=folder_name
        ).first()
        
        if not folder:
            return jsonify({'exists': False})
            
        # 检查文档是否在收藏夹中
        exists = FolderContents.query.filter_by(
            Folder_id=folder.Folder_id,
            Doc_id=doc_id
        ).first() is not None
        
        return jsonify({'exists': exists})
        
    except Exception as e:
        return jsonify({
            'exists': False,
            'error': str(e)
        })

@reader_bp.route('/get_text/<doc_id>/<text_type>')
@login_required
def get_text(doc_id, text_type):
    try:
        doc = validate_document_access(doc_id)
        
        if text_type == 'simplified':
            text = doc.Doc_simplifiedText
        else:
            text = doc.Doc_originalText
            
        if not text:
            return jsonify({
                'success': False,
                'error': '文本不存在'
            })
            
        return jsonify({
            'success': True,
            'text': text
        })
        
    except Exception as e:
        logger.error(f"获取文本失败: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@reader_bp.route('/add_evernote', methods=['POST'])
@login_required
def add_evernote():
    try:
        data = request.get_json()
        
        # 验证必要字段
        if not all(field in data for field in ['doc_id', 'content']):
            return jsonify({
                'success': False,
                'error': '缺少必要字段'
            }), 400
            
        # 创建笔记
        new_note = Evernote(
            Doc_id=data['doc_id'],
            User_id=current_user.User_id,
            Evernote_text=data['content'],
            Evernote_viewedAt=datetime.now()
        )
        
        db.session.add(new_note)
        db.session.commit()
        
        # 只创建一条用户操作的审计日志
        create_audit_log(
            current_user.User_id,
            'INSERT',
            f'用户 {current_user.User_name} 在文档 {data["doc_id"]} 中添加了笔记',
            'Evernote',
            new_note.Evernote_id
        )
        
        return jsonify({
            'success': True,
            'data': {
                'id': new_note.Evernote_id,
                'content': new_note.Evernote_text,
                'created_at': new_note.Evernote_viewedAt.strftime('%Y-%m-%d %H:%M:%S')
            }
        })
        
    except Exception as e:
        logger.error(f"添加笔记时出错: {str(e)}")
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@reader_bp.route('/get_evernotes/<doc_id>')
@login_required
def get_evernotes(doc_id):
    try:
        notes = Evernote.query.filter_by(
            Doc_id=doc_id,
            User_id=current_user.User_id
        ).order_by(Evernote.Evernote_viewedAt.desc()).all()
        
        return jsonify({
            'success': True,
            'notes': [{
                'id': note.Evernote_id,
                'content': note.Evernote_text,
                'created_at': note.Evernote_viewedAt.strftime('%Y-%m-%d %H:%M:%S')
            } for note in notes]
        })
        
    except Exception as e:
        logger.error(f"获取笔记列表失: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@reader_bp.route('/delete_evernote/<int:note_id>', methods=['DELETE'])
@login_required
def delete_evernote(note_id):
    try:
        note = Evernote.query.filter_by(
            Evernote_id=note_id,
            User_id=current_user.User_id
        ).first()
        
        if not note:
            return jsonify({
                'success': False,
                'error': '笔记不存在或无权限删除'
            }), 404
            
        db.session.delete(note)
        db.session.commit()
        
        create_audit_log(
            current_user.User_id,
            'DELETE',
            f'删除笔记 {note_id}',
            'Evernote',
            note_id
        )
        
        return jsonify({
            'success': True,
            'message': '笔记删除成功'
        })
        
    except Exception as e:
        logger.error(f"删除笔记时出错: {str(e)}")
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@reader_bp.route('/update_highlight_color/<int:highlight_id>', methods=['POST'])
@login_required
def update_highlight_color(highlight_id):
    try:
        data = request.get_json()
        highlight = Highlights.query.filter_by(
            Highlight_id=highlight_id,
            User_id=current_user.User_id
        ).first()
        
        if not highlight:
            return jsonify({
                'success': False,
                'error': '高亮不存在或无权限修改'
            }), 404
            
        highlight.Highlight_color = data['color']
        db.session.commit()
        
        return jsonify({
            'success': True
        })
        
    except Exception as e:
        logger.error(f"更新高亮颜色时出错: {str(e)}")
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@reader_bp.route('/update_note/<int:note_id>', methods=['POST'])
@login_required
def update_note(note_id):
    try:
        data = request.get_json()
        note = Notes.query.filter_by(
            Note_id=note_id,
            User_id=current_user.User_id
        ).first()
        
        if not note:
            return jsonify({
                'success': False,
                'error': '批注不存在或无权限修改'
            }), 404
            
        note.Note_annotationText = data['content']
        db.session.commit()
        
        return jsonify({
            'success': True
        })
        
    except Exception as e:
        logger.error(f"更新批注时出错: {str(e)}")
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@reader_bp.route('/update_evernote/<int:note_id>', methods=['PUT'])
@login_required
def update_evernote(note_id):
    try:
        data = request.get_json()
        
        if not data.get('content'):
            return jsonify({
                'success': False,
                'error': '笔记内容不能为空'
            }), 400
            
        note = Evernote.query.filter_by(
            Evernote_id=note_id,
            User_id=current_user.User_id
        ).first()
        
        if not note:
            return jsonify({
                'success': False,
                'error': '笔记不存在或无权限修改'
            }), 404
            
        note.Evernote_text = data['content']
        note.Evernote_viewedAt = datetime.now()
        db.session.commit()
        
        create_audit_log(
            current_user.User_id,
            'UPDATE',
            f'更新笔记 {note_id}',
            'Evernote',
            note_id
        )
        
        return jsonify({
            'success': True,
            'message': '笔记更新成功'
        })
        
    except Exception as e:
        logger.error(f"更新笔记时出错: {str(e)}")
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@reader_bp.route('/delete_comment/<int:comment_id>', methods=['DELETE'])
@login_required
def delete_comment(comment_id):
    try:
        comment = Comments.query.filter_by(
            Comment_id=comment_id,
            User_id=current_user.User_id
        ).first()
        
        if not comment:
            return jsonify({
                'success': False,
                'error': '评论不存在或无权限删除'
            }), 404
            
        db.session.delete(comment)
        db.session.commit()
        
        create_audit_log(
            current_user.User_id,
            'DELETE',
            f'删除评论 {comment_id}',
            'Comments',
            comment_id
        )
        
        return jsonify({
            'success': True,
            'message': '评论删除成功'
        })
        
    except Exception as e:
        logger.error(f"删除评论时出错: {str(e)}")
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@reader_bp.route('/export_document', methods=['POST'], endpoint='export_document_post')
@login_required
def export_document_post():
    try:
        data = request.get_json()
        doc_id = data.get('doc_id')
        
        if not doc_id:
            return jsonify({
                'success': False,
                'error': '文档ID不能为空'
            }), 400
            
        # 验证文档存在
        doc = Documents.query.filter_by(Doc_id=doc_id).first()
        if not doc:
            logger.error(f"找不到文档: {doc_id}")
            return jsonify({
                'success': False,
                'error': f'找不到ID为 {doc_id} 的文档'
            }), 404
            
        format = data.get('format', 'word')  # 默认为word格式
        include_highlights = data.get('highlights', True)
        include_annotations = data.get('annotations', True)
        include_notes = data.get('notes', True)
        include_comments = data.get('comments', True)
        
        # 创建Word文档
        word_doc = Document()
        
        # 添加标题
        word_doc.add_heading(doc.Doc_title or '无标题文档', 0)
        
        # 添加基本信息
        word_doc.add_paragraph(f'文书类型: {doc.Doc_type or "未知类型"}')
        if doc.Doc_summary:
            word_doc.add_paragraph(f'文书大意: {doc.Doc_summary}')
        
        # 添加正文
        word_doc.add_heading('正文', level=1)
        content = doc.Doc_originalText or ''
        if not content:
            word_doc.add_paragraph('文档内容为空')
        else:
            word_doc.add_paragraph(content)
        
        # 添加用户的标注内容
        if current_user.is_authenticated:
            if include_highlights:
                try:
                    highlights = Highlights.query.filter_by(
                        Doc_id=doc_id,
                        User_id=current_user.User_id
                    ).all()
                    
                    if highlights:
                        word_doc.add_heading('高亮部分', level=2)
                        for h in highlights:
                            try:
                                if h.Highlight_startPosition is not None and h.Highlight_endPosition is not None:
                                    highlighted_text = content[h.Highlight_startPosition:h.Highlight_endPosition]
                                    p = word_doc.add_paragraph()
                                    p.add_run(f'• {highlighted_text}')
                                    p.add_run(f' (标注时间: {h.Highlight_createdAt.strftime("%Y-%m-%d %H:%M:%S")})')
                            except Exception as e:
                                logger.error(f"处理高亮文本时出错: {str(e)}")
                except Exception as e:
                    logger.error(f"获取高亮列表时出错: {str(e)}")
            
            if include_annotations:
                try:
                    notes = Notes.query.filter_by(
                        Doc_id=doc_id,
                        User_id=current_user.User_id
                    ).all()
                    
                    if notes:
                        word_doc.add_heading('批注内容', level=2)
                        for n in notes:
                            try:
                                if n.Note_startPosition is not None and n.Note_endPosition is not None:
                                    noted_text = content[n.Note_startPosition:n.Note_endPosition]
                                    p = word_doc.add_paragraph()
                                    p.add_run(f'原文: {noted_text}\n').bold = True
                                    p.add_run(f'批注: {n.Note_annotationText}')
                                    p.add_run(f'\n(批注时间: {n.Note_createdAt.strftime("%Y-%m-%d %H:%M:%S")})')
                            except Exception as e:
                                logger.error(f"处理批注文本时出错: {str(e)}")
                except Exception as e:
                    logger.error(f"获取批注列表时出错: {str(e)}")
            
            if include_notes:
                try:
                    evernotes = Evernote.query.filter_by(
                        Doc_id=doc_id,
                        User_id=current_user.User_id
                    ).all()
                    
                    if evernotes:
                        word_doc.add_heading('笔记内容', level=2)
                        for note in evernotes:
                            p = word_doc.add_paragraph()
                            p.add_run(note.Evernote_text)
                            p.add_run(f'\n(记录时间: {note.Evernote_viewedAt.strftime("%Y-%m-%d %H:%M:%S")})')
                except Exception as e:
                    logger.error(f"获取笔记列表时出错: {str(e)}")
        
        # 添加评论区
        if include_comments:
            try:
                comments = Comments.query.filter_by(Doc_id=doc_id).all()
                if comments:
                    word_doc.add_heading('评论区', level=1)
                    for comment in comments:
                        try:
                            p = word_doc.add_paragraph()
                            p.add_run(f'• {comment.Comment_text}')
                            user_name = comment.user.User_name if comment.user else '匿名用户'
                            p.add_run(f'\n  by {user_name} ({comment.Comment_createdAt.strftime("%Y-%m-%d %H:%M:%S")})')
                        except Exception as e:
                            logger.error(f"处理评论时出错: {str(e)}")
            except Exception as e:
                logger.error(f"获取评论列表时出错: {str(e)}")
        
        # 保存到内存中
        file_stream = BytesIO()
        word_doc.save(file_stream)
        file_stream.seek(0)
        
        # 创建审计日志
        try:
            create_audit_log(
                current_user.User_id,
                'EXPORT',
                f'导出文档 {doc_id}',
                'Documents',
                doc_id
            )
        except Exception as e:
            logger.error(f"创建审计日志时出错: {str(e)}")
        
        # 返回文件
        filename = f"{doc.Doc_title or '导出文档'}.docx"
        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        logger.error(f"导出文档时出错: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'error': f'导出文档失败: {str(e)}'
        }), 500

@reader_bp.route('/call_llm', methods=['POST'])
@login_required
def call_llm():
    try:
        data = request.get_json()
        
        # 验证必要字段
        if not data.get('prompt'):
            return jsonify({
                'success': False,
                'error': '提示词不能为空'
            }), 400
            
        # 获取API配置
        api_key = current_app.config.get('MOONSHOT_API_KEY')
        base_url = current_app.config.get('MOONSHOT_BASE_URL')
        model = current_app.config.get('MOONSHOT_MODEL')
        system_prompt = current_app.config.get('MOONSHOT_SYSTEM_PROMPT')
        
        if not api_key or not base_url:
            return jsonify({
                'success': False,
                'error': 'API配置未设置'
            }), 500
            
        # 创建OpenAI客户端
        client = OpenAI(
            api_key=api_key,
            base_url=base_url
        )
        
        try:
            # 调用API
            completion = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": data['prompt']}
                ],
                temperature=0.3
            )
            
            # 获取响应
            response = completion.choices[0].message.content
            
            # 创建审计日志
            create_audit_log(
                current_user.User_id,
                'LLM_CALL',
                f'用户调用了KIMI大模型API',
                'System',
                None
            )
            
            return jsonify({
                'success': True,
                'response': response
            })
            
        except Exception as e:
            logger.error(f"KIMI API调用失败: {str(e)}")
            return jsonify({
                'success': False,
                'error': f'API调用失败: {str(e)}'
            }), 500
            
    except Exception as e:
        logger.error(f"处理请求时出错: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@reader_bp.route('/get_corrections/<doc_id>')
@login_required
def get_corrections(doc_id):
    try:
        corrections = Corrections.query.filter_by(Doc_id=doc_id).all()
        
        return jsonify({
            'success': True,
            'corrections': [{
                'correction_id': c.Correction_id,
                'user_id': str(c.User_id),  # 转换为字符串以便于前端比较
                'user_name': Users.query.get(c.User_id).User_name if c.User_id else '系统',
                'text': c.Correction_text,
                'created_at': c.Correction_createdAt.strftime('%Y-%m-%d %H:%M:%S')
            } for c in corrections]
        })
        
    except Exception as e:
        logger.error(f"获取纠错记录失败: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@reader_bp.route('/add_correction', methods=['POST'])
@login_required
def add_correction():
    try:
        data = request.get_json()
        
        # 验证必要字段
        if not data.get('doc_id') or not data.get('text'):
            return jsonify({
                'success': False,
                'error': '缺少必要字段'
            }), 400
            
        # 创建纠错记录
        correction = Corrections(
            Doc_id=data['doc_id'],
            User_id=current_user.User_id,
            Correction_text=data['text'],
            Correction_createdAt=datetime.now()
        )
        
        db.session.add(correction)
        db.session.commit()
        
        # 只创建一条用户操作的审计日志
        create_audit_log(
            current_user.User_id,
            'INSERT',
            f'用户 {current_user.User_name} 在文档 {data["doc_id"]} 中添加了纠错',
            'Corrections',
            correction.Correction_id
        )
        
        return jsonify({
            'success': True,
            'data': {
                'correction_id': correction.Correction_id,
                'user_id': str(correction.User_id),
                'user_name': current_user.User_name,
                'text': correction.Correction_text,
                'created_at': correction.Correction_createdAt.strftime('%Y-%m-%d %H:%M:%S')
            }
        })
        
    except Exception as e:
        logger.error(f"添加纠错失败: {str(e)}")
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@reader_bp.route('/delete_correction/<int:correction_id>', methods=['DELETE'])
@login_required
def delete_correction(correction_id):
    try:
        correction = Corrections.query.filter_by(
            Correction_id=correction_id,
            User_id=current_user.User_id
        ).first()
        
        if not correction:
            return jsonify({
                'success': False,
                'error': '纠错记录不存在或无权限删除'
            }), 404
            
        db.session.delete(correction)
        db.session.commit()
        
        # 创建审计日志
        create_audit_log(
            current_user.User_id,
            'DELETE',
            f'用户删除了文档 {correction.Doc_id} 的纠错',
            'Corrections',
            correction_id
        )
        
        return jsonify({
            'success': True,
            'message': '纠错记录删除成功'
        })
        
    except Exception as e:
        logger.error(f"删除纠错失败: {str(e)}")
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500
    
@reader_bp.route('/get_folders', methods=['GET'])
@login_required
def get_folders():
    """获取用户的收藏夹列表"""
    try:
        # 获取当前用户的所有收藏夹
        folders = Folders.query.filter_by(User_id=current_user.User_id).all()
        
        # 将收藏夹信息转换为JSON格式
        folders_list = [{'name': folder.Folder_name, 'id': folder.Folder_id} for folder in folders]
        
        return jsonify({
            'success': True,
            'folders': folders_list
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500
